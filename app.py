from flask import Flask, render_template, request, session, jsonify
from flask import send_file, session, request, flash
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from flask import Flask, render_template, request, jsonify, redirect, url_for, session
from flask_cors import CORS
import mysql.connector
import random
import re
from pythainlp.tokenize import word_tokenize
from pythainlp.tag import pos_tag
import pyrebase

from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from flask import send_file
from datetime import datetime

import io

"""
credentialsFile = "mathprob-cfefb-firebase-adminsdk-ul8au-34bdcf751f.json"
jsonFile = "client_secret_550663957778-bfphkjgnnma9cskqm3bl3o9bnhat72j4.apps.googleusercontent.com.json"
SCOPES = ["https://www.googleapis.com/auth/cloud-platform", "https://www.googleapis.com/auth/calendar.readonly"]
from google.oauth2 import service_account
credentials = service_account.Credentials.from_service_account_file(credentialsFile, scopes=SCOPES)
# scoped_credentials = credentials.with_scopes(SCOPES)
"""

app = Flask(__name__)
CORS(app)
app.secret_key = 'your_secret_key'  # ใช้ secret key สำหรับการเข้ารหัส session

"""
const firebaseConfig = {
  apiKey: "AIzaSyANjn9wPeoV07Y7XxuGKDIoK8yerNN9LQI",
  authDomain: "mathprob-cfefb.firebaseapp.com",
  databaseURL: "https://mathprob-cfefb-default-rtdb.asia-southeast1.firebasedatabase.app",
  projectId: "mathprob-cfefb",
  storageBucket: "mathprob-cfefb.appspot.com",
  messagingSenderId: "550663957778",
  appId: "1:550663957778:web:20059a516f561c8c01d7cb",
  measurementId: "G-BGEWQB60Y3"
};
"""

# ตั้งค่า firebase
config = {
    "apiKey": "AIzaSyANjn9wPeoV07Y7XxuGKDIoK8yerNN9LQI",
    "authDomain": "mathprob-cfefb.firebaseapp.com",
    "databaseURL": "https://mathprob-cfefb-default-rtdb.asia-southeast1.firebasedatabase.app",
    "storageBucket": "mathprob-cfefb.appspot.com"
}
# Initialize Firebase
firebase = pyrebase.initialize_app(config)
# Get reference to the auth service and database service
auth = firebase.auth()
dbFirebase = firebase.database()

# เชื่อมต่อMySQL
host = "localhost"
user = "root"
password = ""
db = "mathprob"
db2 = "word"


# Utility functions for fetching nouns, verbs, and objects


def get_random_noun(math_symbol=None):
    # ถ้าเป็นการหาร ให้ดึงคำนามเฉพาะสำหรับโจทย์การหาร
    if math_symbol == '/':
        return fetch_random_word("word_nouns1", "nouns", db2, where_column="type", where_value="divide")
    else:
        # ใช้คำนามจากฐานข้อมูลทุกคำ สำหรับสัญลักษณ์อื่นๆ
        return fetch_random_word("word_nouns1", "nouns", db2)


def get_random_object(problem_type):
    return fetch_random_word("word_object", "nouns", db2, "type", problem_type)


def get_random_verb(math_symbol):
    return fetch_random_word("word_verb", "verb", db2, "mathsym", math_symbol)


def fetch_random_word(table, column, database, where_column=None, where_value=None, additional_condition=None):
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=database)
        mycursor = mydb.cursor(dictionary=True)

        query = f"SELECT {column} FROM {table}"

        conditions = []
        if where_column and where_value:
            conditions.append(f"{where_column} = %s")
        if additional_condition:
            conditions.append(f"{additional_condition[0]} = %s")

        if conditions:
            query += " WHERE " + " AND ".join(conditions)

        params = []
        if where_value:
            params.append(where_value)
        if additional_condition:
            params.append(additional_condition[1])

        print(f"Executing query: {query} with params: {params}")  # เพิ่ม debug
        mycursor.execute(query, tuple(params))
        words = mycursor.fetchall()

        return random.choice(words)[column] if words else None
    except mysql.connector.Error as err:
        print(f"Error: {err}")
        return None
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()


def get_all_words(table, column, database, where_column=None, where_value=None):
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=database)
        mycursor = mydb.cursor(dictionary=True)

        query = f"SELECT {column} FROM {table}"

        # เพิ่มเงื่อนไข (ถ้ามี)
        if where_column and where_value:
            query += f" WHERE {where_column} = %s"
            mycursor.execute(query, (where_value,))
        else:
            mycursor.execute(query)

        words = mycursor.fetchall()
        return [word[column] for word in words] if words else []
    except mysql.connector.Error as err:
        print(f"Error: {err}")
        return []
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()


def get_operation_type(symbol):
    if symbol == '+':
        return 'add'
    elif symbol == '-':
        return 'minus'
    elif symbol == '*':
        return 'multiply'
    elif symbol == '/':
        return 'divide'
    return None


def get_random_pattern(problem_type, math_symbol):
    # แปลงสัญลักษณ์เป็นคำที่ตรงกับ mathsym
    operation_type = get_operation_type(math_symbol)
    if operation_type is None:
        return None  # ถ้าไม่มี operation_type ที่ตรงกันก็คืนค่า None

    return fetch_random_word("mathprob_pattern", "pattern", db, where_column="type", where_value=problem_type, additional_condition=("mathsym", operation_type))


def get_random_classifier(problem_type, math_symbol):
    # แปลงสัญลักษณ์ทางคณิตศาสตร์เป็นคำที่ใช้ในฐานข้อมูล
    db_math_symbol = 'multiply' if math_symbol == '*' else 'divide' if math_symbol == '/' else None
    if db_math_symbol is None:
        return None

    # ดึงลักษณะนามที่ตรงกับประเภทของโจทย์และสัญลักษณ์ทางคณิตศาสตร์
    return fetch_random_word("word_classifier", "classifier", db2, where_column="type", where_value=problem_type, additional_condition=("mathsym", db_math_symbol))


def generate_problem_with_multiplication_pattern(math_symbol, digit_choice, problem_type):
    num1, num2 = generate_random_number_by_digits(
        digit_choice), generate_random_number_by_digits(digit_choice)

    # ดึง pattern สำหรับโจทย์การคูณ
    pattern = get_random_pattern(problem_type, math_symbol)
    if not pattern:
        return None

    # แทนค่า X และ Y ใน pattern ด้วยจำนวนตัวเลข
    pattern = pattern.replace(
        "X", f'<span contenteditable="false">{num1:,}</span>')
    pattern = pattern.replace(
        "Y", f'<span contenteditable="false">{num1:,}</span>')

    # ดึงคำนาม, กรรม, และลักษณะนาม
    subject = get_random_noun(math_symbol)
    obj = get_random_object(problem_type)
    classifier = get_random_classifier
    subject2 = get_random_noun(math_symbol)

    # สุ่มคำนามตัวที่สองให้ไม่ซ้ำกับตัวแรก
    subject2 = get_random_noun(math_symbol)
    while subject2 == subject:
        subject2 = get_random_noun()

    # ทำการแทนที่ใน pattern ด้วย <span> เพื่อให้แก้ไขได้
    if subject:
        pattern = pattern.replace(
            "(คำนาม)", f'<span class="editable-word" data-type="noun" contenteditable="true">{subject}</span>', 1)
    if obj:
        pattern = pattern.replace(
            "(กรรม)", f'<span class="editable-word" data-type="object" contenteditable="true">{obj}</span>')
    if classifier:
        pattern = pattern.replace(
            "(ลักษณะนาม)", f'<span class="editable-word" data-type="classifier" contenteditable="true">{classifier}</span>')
    if subject2:
        pattern = pattern.replace(
            "(คำนาม)", f'<span class="editable-word" data-type="noun" contenteditable="true">{subject2}</span>', 1)

    # เพิ่มการตัดคำเพื่อหา "หน่วย" จากโจทย์ที่สร้างขึ้นมา
    tokenized = word_tokenize(pattern)
    unit = tokenized[-2]  # คำสุดท้ายของโจทย์

    answer = calculate_answer(num1, num2, math_symbol)

    return pattern, answer, unit


def generate_problem_with_division_pattern(math_symbol, digit_choice, problem_type):
    # สร้างตัวหาร (num2) ที่ไม่เป็นศูนย์
    num2 = generate_random_number_by_digits(digit_choice)
    while num2 == 0:
        num2 = generate_random_number_by_digits(digit_choice)

    # สร้างตัวตั้ง (num1) โดยทำให้แน่ใจว่ามันหารด้วย num2 ลงตัว
    multiplier = random.randint(1, 10)  # เลือกตัวคูณสุ่มเพื่อสร้าง num1
    num1 = num2 * multiplier

    # ดึง pattern สำหรับโจทย์การหาร
    pattern = get_random_pattern(problem_type, math_symbol)
    if not pattern:
        return None

    # แทนที่ X, Y ใน pattern ด้วยตัวเลขสุ่มและใช้ `<span>` สำหรับให้แก้ไขได้
    pattern = pattern.replace(
        "X", f'<span contenteditable="false">{num1:,}</span>')
    pattern = pattern.replace(
        "Y", f'<span contenteditable="false">{num1:,}</span>')

    # ดึงข้อมูลเพื่อใช้แทนใน pattern
    subject = get_random_noun(math_symbol)
    obj = get_random_object(problem_type)
    classifier = get_random_classifier(problem_type, math_symbol)
    verb = get_random_verb(math_symbol)
    subject2 = get_random_noun(math_symbol)

    # สุ่มคำนามตัวที่สองให้ไม่ซ้ำกับตัวแรก
    while subject2 == subject:
        subject2 = get_random_noun(math_symbol)

    # ทำการแทนที่ใน pattern ด้วย `<span>` เพื่อให้แก้ไขได้
    if subject:
        pattern = pattern.replace(
            "(คำนาม)", f'<span class="editable-word" data-type="noun" contenteditable="true">{subject}</span>', 1)
    if obj:
        pattern = pattern.replace(
            "(กรรม)", f'<span class="editable-word" data-type="object" contenteditable="true">{obj}</span>')
    if verb:
        pattern = pattern.replace(
            "(คำกริยา)", f'<span class="editable-word" data-type="verb" contenteditable="true">{verb}</span>')
    if classifier:
        pattern = pattern.replace(
            "(ลักษณะนาม)", f'<span class="editable-word" data-type="classifier" contenteditable="true">{classifier}</span>')
    if subject2:
        pattern = pattern.replace(
            "(คำนาม)", f'<span class="editable-word" data-type="noun" contenteditable="true">{subject2}</span>', 1)

    answer = calculate_answer(num1, num2, math_symbol)

    # เพิ่มการตัดคำเพื่อหา "หน่วย" จากโจทย์ที่สร้างขึ้นมา
    tokenized = word_tokenize(pattern)
    unit = tokenized[-2]  # คำสุดท้ายของโจทย์
    return pattern, answer, unit


# Function to calculate answer based on math_symbol

def calculate_answer(num1, num2, math_symbol):
    try:
        if math_symbol == '+':
            answer = num1 + num2
        elif math_symbol == '-':
            answer = num1 - num2
        elif math_symbol == '*':
            answer = num1 * num2
        elif math_symbol == '/':
            answer = num1 / num2 if num2 != 0 else 'Cannot divide by zero'
        else:
            return None

        # เพิ่มลูกน้ำในคำตอบที่เป็นตัวเลขจำนวนเต็ม
        if isinstance(answer, (int, float)) and answer == int(answer):
            return f"{int(answer):,}"
        elif isinstance(answer, float):
            return f"{answer:,.2f}"  # แสดงทศนิยม 2 ตำแหน่ง
        return answer
    except ZeroDivisionError:
        return 'Cannot divide by zero'

# Generate problem with pattern


def get_random_verb(math_symbol):
    # แปลง math_symbol ให้ตรงกับค่าที่เก็บในฐานข้อมูล
    db_math_symbol = symbol_mapping.get(math_symbol, math_symbol)
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2)
        mycursor = mydb.cursor(dictionary=True)

        # ตรวจสอบการทำงานของ query
        query = "SELECT verb FROM word_verb WHERE mathsym = %s"
        print(f"Executing query: {query} with math_symbol = {db_math_symbol}")
        mycursor.execute(query, (db_math_symbol,))
        verbs = mycursor.fetchall()

        # ตรวจสอบว่ามีคำที่ดึงออกมาได้หรือไม่
        if verbs:
            selected_word = random.choice(verbs)['verb']
            print(f"Selected verb: {selected_word}")
            return selected_word
        else:
            print(f"No verb found for math_symbol: {db_math_symbol}")
            return None
    except mysql.connector.Error as err:
        print(f"Error: {err}")
        return None
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()


# ในการตรวจสอบกรณีการสร้างโจทย์แบบสุ่ม

def generate_problem_with_pattern(math_symbol, digit_choice, problem_type):
    used_numbers = set()  # ใช้ set เก็บตัวเลขที่เคยใช้แล้ว
    unique_attempts = 0  # ใช้เพื่อตรวจสอบจำนวนครั้งที่พยายามหาตัวเลขไม่ซ้ำกัน

    def get_unique_number(digit_choice):
        nonlocal unique_attempts
        while True:
            num = generate_random_number_by_digits(digit_choice)
            if num not in used_numbers or unique_attempts > 5:  # เพิ่มเงื่อนไขเพื่อหลุด loop
                used_numbers.add(num)
                return num
            unique_attempts += 1

    if math_symbol == '/':
        # สร้างตัวหาร (num2) ที่ไม่เป็นศูนย์
        num2 = get_unique_number(digit_choice)
        while num2 == 0:
            num2 = get_unique_number(digit_choice)

        # สร้างตัวตั้ง (num1) โดยทำให้แน่ใจว่ามันหารด้วย num2 ลงตัว
        multiplier = random.randint(1, 10)  # เลือกตัวคูณสุ่มเพื่อสร้าง num1
        num1 = num2 * multiplier
    else:
        # สำหรับการบวก ลบ หรือ คูณ ให้สุ่มตัวเลขปกติ
        num1 = get_unique_number(digit_choice)
        num2 = get_unique_number(digit_choice)

    pattern = get_random_pattern(problem_type, math_symbol)
    if not pattern:
        return None

    # แทนที่ X, Y ใน pattern ด้วยตัวเลขสุ่มและใช้ `<span>` สำหรับให้แก้ไขได้
    pattern = pattern.replace(
        "X", f'<span contenteditable="false">{num1:,}</span>')  # แสดงตัวเลขพร้อมลูกน้ำ
    pattern = pattern.replace(
        "Y", f'<span contenteditable="false">{num2:,}</span>')  # แสดงตัวเลขพร้อมลูกน้ำ

    # ดึงข้อมูลเพื่อใช้แทนใน pattern
    subject = get_random_noun(math_symbol)
    obj = get_random_object(problem_type)
    verb = get_random_verb(math_symbol)
    classifier = get_random_classifier(problem_type, math_symbol) if math_symbol in [
        '*', '/'] else None
    subject2 = get_random_noun()

    # สุ่มคำนามตัวที่สองให้ไม่ซ้ำกับตัวแรก
    while subject2 == subject:
        subject2 = get_random_noun(math_symbol)

    # ทำการแทนที่ใน pattern ด้วย `<span>` เพื่อให้แก้ไขได้
    if subject:
        pattern = pattern.replace(
            "(คำนาม)", f'<span class="editable-word" data-type="noun" contenteditable="true">{subject}</span>', 1)
    if obj:
        pattern = pattern.replace(
            "(กรรม)", f'<span class="editable-word" data-type="object" contenteditable="true">{obj}</span>')
    if verb:
        pattern = pattern.replace(
            "(คำกริยา)", f'<span class="editable-word" data-type="verb" contenteditable="true">{verb}</span>')
    if classifier:
        pattern = pattern.replace(
            "(ลักษณะนาม)", f'<span class="editable-word" data-type="classifier" contenteditable="true">{classifier}</span>')
    if subject2:
        pattern = pattern.replace(
            "(คำนาม)", f'<span class="editable-word" data-type="noun" contenteditable="true">{subject2}</span>', 1)

    # เพิ่มการตัดคำเพื่อหา "หน่วย" จากโจทย์ที่สร้างขึ้นมา
    tokenized = word_tokenize(pattern)
    unit = tokenized[-2]  # คำสุดท้ายของโจทย์

    answer = calculate_answer(num1, num2, math_symbol)
    return pattern, answer, unit


# แยกตัวเลขและสัญลักษณ์ทางคณิตศาสตร์ออกจากประโยค


def extractmathsym(sentence):
    pattern = re.compile(r'\d+')
    matches = pattern.findall(sentence)
    return matches if matches else []


def generate_random_number_by_digits(digit_type):
    if digit_type == "หน่วย":
        return random.randint(1, 9)
    elif digit_type == "สิบ":
        return random.randint(10, 99)
    elif digit_type == "ร้อย":
        return random.randint(100, 999)
    elif digit_type == "พัน":
        return random.randint(1000, 9999)
    return 0


# Mapping สัญลักษณ์ไปยังค่าที่ใช้ในฐานข้อมูล
symbol_mapping = {
    '+': 'add',
    '-': 'minus',
    '*': 'multiply',
    '/': 'divide'
}


def generate_problem_with_pattern_from_sentence(symsentence, problem_type):
    # แยกตัวเลขออกจากประโยคสัญลักษณ์
    numbers = extractmathsym(symsentence)
    if len(numbers) < 2:
        return None, None, None, None

    # แปลงตัวเลขในประโยคสัญลักษณ์
    num1, num2 = int(numbers[0]), int(numbers[1])
    # กำหนด math_symbol ตามเครื่องหมายในประโยคสัญลักษณ์
    math_symbol = '+' if '+' in symsentence else '-' if '-' in symsentence else '*' if '*' in symsentence else '/'

    # ตรวจสอบหากเป็นการหารและตัวหารเป็นศูนย์
    if math_symbol == '/' and num2 == 0:
        return None, None, None, None

    # ดึง pattern ตามประเภทและสัญลักษณ์คณิตศาสตร์
    pattern = get_random_pattern(problem_type, math_symbol)
    if not pattern:
        return None, None, None, None

    # แทนที่ X และ Y ใน pattern ด้วยตัวเลขที่ได้
    pattern = pattern.replace(
        "X", f'<span contenteditable="false">{num1}</span>')
    pattern = pattern.replace(
        "Y", f'<span contenteditable="false">{num2}</span>')

    # ดึงข้อมูลจากฐานข้อมูลหรือคำเพื่อใช้ในโจทย์
    subject = get_random_noun(math_symbol)
    obj = get_random_object(problem_type)
    verb = get_random_verb(math_symbol)
    classifier = get_random_classifier(
        problem_type, math_symbol) if math_symbol in ['*', '/'] else None
    subject2 = get_random_noun()

    # สุ่มคำนามตัวที่สองให้ไม่ซ้ำกับตัวแรก
    while subject2 == subject:
        subject2 = get_random_noun(math_symbol)

    # ทำการแทนที่ใน pattern ด้วย `<span>` เพื่อให้สามารถแก้ไขได้
    if subject:
        pattern = pattern.replace(
            "(คำนาม)", f'<span class="editable-word" data-type="noun" contenteditable="true">{subject}</span>', 1)
    if obj:
        pattern = pattern.replace(
            "(กรรม)", f'<span class="editable-word" data-type="object" contenteditable="true">{obj}</span>')
    if verb:
        pattern = pattern.replace(
            "(คำกริยา)", f'<span class="editable-word" data-type="verb" contenteditable="true">{verb}</span>')
    if classifier:
        pattern = pattern.replace(
            "(ลักษณะนาม)", f'<span class="editable-word" data-type="classifier" contenteditable="true">{classifier}</span>')
    if subject2:
        pattern = pattern.replace(
            "(คำนาม)", f'<span class="editable-word" data-type="noun" contenteditable="true">{subject2}</span>', 1)

    # แยกคำจาก pattern เพื่อดึงคำสุดท้ายที่เป็นหน่วยของโจทย์
    tokenized = word_tokenize(pattern)
    unit = tokenized[-2] if len(tokenized) > 1 else ''

    # คำนวณคำตอบจาก num1 และ num2
    answer = calculate_answer(num1, num2, math_symbol)

    # คืนค่าผลลัพธ์ทั้งหมดพร้อม math_symbol สำหรับใช้ในอนาคต
    return pattern, answer, unit, math_symbol


@app.route('/home', methods=['POST', 'GET'])
def index():
    # ดึงข้อมูลจาก session หรือสร้างใหม่หากยังไม่มี
    updated_problems = session.get('updated_problems', [])
    answers = session.get('answers', [])
    units = session.get('units', [])
    problem_types = session.get('problem_types', [])
    error = None

    if request.method == 'POST':
        # รีเซ็ตค่า parameters ใหม่ก่อนการสร้างโจทย์ทุกครั้ง
        parameters = []  # ล้างค่า parameters ให้เป็นค่าว่าง
        form = request.form

        # ตรวจสอบกรณีการสร้างโจทย์แบบสุ่ม
        if 'math_symbol' in form and 'digit_choice' in form and 'noun_category' in form:
            math_symbol = form['math_symbol']
            digit_choice = form['digit_choice']
            problem_type = form['noun_category']
            problem_count = int(form.get('problem_count', 1))

            # กำหนดรายการประเภทที่สามารถสุ่มได้
            available_types = ['object', 'fruit',
                               'vehicle', 'animal', 'plant', 'food']

            # สร้างโจทย์และเก็บใน session
            for _ in range(problem_count):
                # หากเลือก "สุ่มประเภท" ให้สุ่มประเภทใหม่สำหรับแต่ละโจทย์
                current_type = random.choice(
                    available_types) if problem_type == 'random' else problem_type

                # สร้างโจทย์โดยใช้ current_type
                generated_problem, answer, unit = generate_problem_with_pattern(
                    math_symbol, digit_choice, current_type)

                if generated_problem:
                    updated_problems.append(generated_problem)
                    answers.append(answer)
                    units.append(unit)
                    # ใช้ current_type แทน 'random'
                    problem_types.append(current_type)
                    parameters.append({
                        'math_symbol': math_symbol,
                        'digit_choice': digit_choice,
                        'noun_category': current_type
                    })

                    print(f"Math Symbol saved in parameters: {math_symbol}")
                    print(f"Current parameters list: {parameters}")

                    print("Created problem with parameters:")
                    print(
                        f"Math Symbol: {math_symbol}, Digit Choice: {digit_choice}, Noun Category: {current_type}")
                    print(
                        f"Generated Problem: {generated_problem}, Answer: {answer}, Unit: {unit}")

                else:
                    error = "No pattern found for the given type and symbol."
                    break

        elif 'symsen' in form and 'noun_category' in form:
            symsentences = form['symsen']
            problem_type = form['noun_category']

            # สุ่มประเภทในกรณีที่เลือกเป็น "random"
            available_types = ['object', 'fruit', 'vehicle', 'animal', 'plant', 'food']

            sentence_list = symsentences.split(',')

            for symsentence in sentence_list:
                symsentence = symsentence.strip()

                # สุ่มประเภทใหม่สำหรับแต่ละประโยคถ้าเลือก "random"
                current_type = random.choice(
                    available_types) if problem_type == 'random' else problem_type

                generated_problem, answer, unit, math_symbol = generate_problem_with_pattern_from_sentence(
                    symsentence, current_type)

                if generated_problem:
                    updated_problems.append(generated_problem)
                    answers.append(answer)
                    units.append(unit)
                    problem_types.append('symsentence')
                    parameters.append({
                        'symsentence': symsentence,
                        'noun_category': current_type,  # ใช้ current_type แทน problem_type
                        'math_symbol': math_symbol  # บันทึก math_symbol ด้วย
                    })

                    print("Created symbolic sentence problem with parameters:")
                    print(
                        f"Symbolic Sentence: {symsentence}, Noun Category: {current_type}")
                    print(
                        f"Generated Problem: {generated_problem}, Answer: {answer}, Unit: {unit}")

                else:
                    error = "Invalid symbol sentence format or no matching pattern."
                    break


        print("Final parameters list being saved to session:", parameters)
        session['parameters'] = parameters

    # เก็บข้อมูลใน session ไม่ให้รีเซ็ตค่าทุกครั้งที่โหลดหน้า
    session['updated_problems'] = updated_problems
    session['answers'] = answers
    session['units'] = units
    session['problem_types'] = problem_types

    return render_template(
        'index.html',
        updated_problems=updated_problems,
        answers=answers,
        units=units,
        problem_types=problem_types,
        parameters=session.get('parameters', []),
        error=error
    )



@app.route('/regen_single_problem', methods=['POST'])
def regen_single_problem():
    data = request.get_json()
    print("Received data for regen:", data)

    try:
        # ตรวจสอบค่าของ index ที่ได้รับจาก client
        index = data.get('index')
        if index is None or not isinstance(index, int) or index >= len(session.get('updated_problems', [])):
            return jsonify({'error': 'Invalid index provided'}), 400

        # ดึงข้อมูลโจทย์จาก session ตาม index
        parameters = session.get('parameters', [])
        if index >= len(parameters):
            return jsonify({'error': 'Invalid index in session parameters'}), 400

        # ใช้ parameters จาก session โดยตรง
        problem_data = parameters[index]
        math_symbol = problem_data.get('math_symbol', '')
        digit_choice = problem_data.get('digit_choice', '')
        noun_category = problem_data['noun_category']
        problem_type = session['problem_types'][index]

        print("Math Symbol being used for regen from session:", math_symbol)
        print("Using parameters:", problem_data)

        # ดำเนินการรีเจนตามประเภท
        if problem_type == 'symsentence':
            symsentence = problem_data.get('symsentence', '')
            print("Symbolic sentence for regen:", symsentence)
            if not symsentence:
                return jsonify({'error': 'Empty symbolic sentence'}), 400

            # เรียกใช้ฟังก์ชัน generate_problem_with_pattern_from_sentence พร้อมรับ math_symbol กลับมา
            result = generate_problem_with_pattern_from_sentence(
                symsentence, noun_category)
            if result is None:
                print("Failed to generate a problem from symbolic sentence.")
                return jsonify({'error': 'Failed to generate a problem from symbolic sentence.'}), 400

            # แยกค่าจากผลลัพธ์
            new_problem, answer, unit, math_symbol = result  # เพิ่ม math_symbol

            # อัปเดต math_symbol ใน parameters เพื่อการรีเจนครั้งต่อไป
            parameters[index]['math_symbol'] = math_symbol
        else:
            # ใช้ math_symbol เฉพาะในกรณีที่ไม่ใช่ symsentence เท่านั้น
            if math_symbol == '/':
                new_problem, answer, unit = generate_problem_with_division_pattern(
                    math_symbol, digit_choice, noun_category)
            elif math_symbol == '*':
                new_problem, answer, unit = generate_problem_with_multiplication_pattern(
                    math_symbol, digit_choice, noun_category)
            else:
                new_problem, answer, unit = generate_problem_with_pattern(
                    math_symbol, digit_choice, noun_category)

        # ตรวจสอบว่าการสร้างโจทย์ใหม่สำเร็จหรือไม่ และอัปเดตข้อมูลใน session
        if new_problem:
            session['updated_problems'][index] = new_problem
            session['answers'][index] = answer
            session['units'][index] = unit
            # บันทึก parameters ที่อัปเดตใหม่
            session['parameters'] = parameters
            session.modified = True  # บันทึกการเปลี่ยนแปลงใน session

            return jsonify({
                'updated_problem': new_problem,
                'answer': answer,
                'unit': unit
            })
        else:
            print("Failed to regenerate the problem.")
            return jsonify({'error': 'Failed to regenerate the problem.'}), 400

    except Exception as e:
        print("An error occurred during regeneration:", e)
        return jsonify({'error': 'An unexpected error occurred'}), 500


@app.route('/delete', methods=['POST'])
def delete():

    problem_to_delete = request.form.get('problem_to_delete')
    updated_problems = session.get('updated_problems', [])
    answers = session.get('answers', [])
    units = session.get('units', [])

    if problem_to_delete and problem_to_delete in updated_problems:
        index_to_delete = updated_problems.index(problem_to_delete)

        if index_to_delete < len(updated_problems):
            updated_problems.pop(index_to_delete)

        if index_to_delete < len(answers):
            answers.pop(index_to_delete)

        if index_to_delete < len(units):
            units.pop(index_to_delete)

        # บันทึกค่าใหม่ลงใน session
        session['updated_problems'] = updated_problems
        session['answers'] = answers
        session['units'] = units

    return redirect(url_for('index'))

# ส่วนของการบันทึกโจทย์ลง firebase


@app.route('/save_all', methods=['POST'])
def save_all():
    # CHECK LOGIN
    if not session.get("is_logged_in"):
        return redirect(url_for('login'))

    updated_problems = session.get('updated_problems', [])
    answers = session.get('answers', [])
    units = session.get('units', [])
    uid = session["uid"]

    # ดึงข้อมูลชุดโจทย์ที่มีอยู่ใน Firebase
    data = dbFirebase.child("problems").child(uid).get().val()

    # ตรวจสอบว่ามีคีย์ 'sets' ใน data หรือไม่ ถ้าไม่มีก็สร้างขึ้นมา
    if not data:
        data = {"sets": []}
    elif 'sets' not in data:
        data['sets'] = []

    # เพิ่มโจทย์ชุดใหม่ลงใน Firebase
    new_set = {
        "problems": updated_problems,
        "answers": answers,
        "units": units,
        "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S")  # เพิ่มวันที่บันทึก
    }
    data["sets"].append(new_set)

    # พิมพ์ข้อมูลเพื่อตรวจสอบว่า new_set ถูกสร้างอย่างถูกต้อง
    print("บันทึกชุดใหม่:", new_set)

    # บันทึกข้อมูลลง Firebase
    dbFirebase.child("problems").child(uid).set(data)

    # ล้างโจทย์จาก session
    session['updated_problems'], session['answers'], session['units'] = [], [], []

    return redirect(url_for('show_all'))

# เเสดงประวัติการสร้างโจทย์


@app.route('/show_all', methods=['GET'])
def show_all():
    # CHECK LOGIN
    if not session.get("is_logged_in"):
        return redirect(url_for('login'))

    # ดึงข้อมูลจาก Firebase
    uid = session["uid"]
    data = dbFirebase.child("problems").child(uid).get().val()

    # ตรวจสอบว่ามีชุดโจทย์อยู่ใน data หรือไม่
    if not data or 'sets' not in data or len(data['sets']) == 0:
        return render_template('show.html', message="ยังไม่มีชุดโจทย์ที่บันทึก")

    # ส่งข้อมูลไปยังหน้าแสดงผลพร้อมกับ enumerate
    return render_template('show.html', data=data, enumerate=enumerate)


# ลบโจทย์ที่เคยบันทึกใน firebase


@app.route('/delete_save', methods=['GET'])
def delete_save():
    uid = session["uid"]
    id = int(request.args.get('id'))
    print("uid:", uid)
    print("id:", id)
    # get old
    data = dbFirebase.child("problems").child(session["uid"]).get().val()
    updated_problems = data["updated_problems"]
    answers = data["answers"]
    units = data["units"]
    del updated_problems[id]
    del answers[id]
    del units[id]
    # update data
    data = {"updated_problems": updated_problems,
            "answers": answers, "units": units}
    dbFirebase.child("problems").child(session["uid"]).set(data)
    return redirect(url_for('show_all'))


@app.route('/clear_all', methods=['POST'])
def clear_all():

    # ลบข้อมูลทั้งหมดใน session
    session['updated_problems'] = []  # กำหนดเป็นลิสต์ว่าง
    session['answers'] = []  # กำหนดเป็นลิสต์ว่าง
    session['units'] = []  # กำหนดเป็นลิสต์ว่าง

    session['select_subject'] = []
    session['select_object'] = []

    return redirect(url_for('index'))


# การบันทึกโจทย์


@app.route('/download_pdf', methods=['POST'])
def download_pdf():
    updated_problems = session.get('updated_problems', [])

    # ตรวจสอบว่าโจทย์ไม่ว่างเปล่า
    if not updated_problems:
        return "No problems to download", 400

    # ลบแท็ก HTML ด้วย BeautifulSoup และจัดการช่องว่างที่ไม่จำเป็น
    cleaned_problems = []
    for problem in updated_problems:
        # ลบแท็ก HTML ทั้งหมดออกด้วย BeautifulSoup
        soup = BeautifulSoup(problem, "html.parser")
        cleaned_text = soup.get_text()  # ดึงข้อความโดยไม่รวมแท็ก
        # ลบช่องว่างเกินที่อาจหลงเหลืออยู่
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        cleaned_problems.append(cleaned_text)

    # สร้างไฟล์ PDF ใน memory
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    # เพิ่มฟอนต์ภาษาไทย
    pdfmetrics.registerFont(TTFont('THSarabun', 'THSarabunNew.ttf'))

    # ตั้งค่าฟอนต์ภาษาไทยและขนาดฟอนต์
    text_object = c.beginText(50, 750)
    text_object.setFont("THSarabun", 16)

    # สร้างหัวกระดาษ ชื่อ,นามสกุล,เลขที่,วันที่
    text_object.textLine(
        f"ชื่อ: ____________  นามสกุล: ____________  เลขที่: ______        วันที่: ____________")
    c.drawText(text_object)

    # ช่องคะแนน
    c.setFont("THSarabun", 16)
    c.drawRightString(580, 750, "คะแนน: ______")

    # สร้างหัวกระดาษ
    c.setFont("THSarabun", 20)
    c.drawCentredString(300, 720, "โจทย์ปัญหาคณิตศาสตร์")

    # ปรับส่วนของโจทย์
    text_object = c.beginText(50, 680)
    text_object.setFont("THSarabun", 16)

    for i, problem in enumerate(cleaned_problems, 1):
        text_object.textLine(f"{i}) {problem}")
        text_object.moveCursor(0, 10)
        # เพิ่มบรรทัดคำว่า "ตอบ" ใต้โจทย์
        text_object.textLine(
            "ตอบ: .....................................................................................")
        text_object.moveCursor(0, 5)

        if text_object.getY() < 100:
            c.drawText(text_object)
            c.showPage()
            text_object = c.beginText(50, 750)
            text_object.setFont("THSarabun", 16)

    c.drawText(text_object)

    # เพิ่มเลขหน้าที่มุมล่างขวา
    for page_num in range(1, c.getPageNumber() + 1):
        c.setFont("THSarabun", 12)
        c.drawRightString(580, 20, f"Page {page_num}")

    c.save()

    # ส่งไฟล์ PDF กลับให้ผู้ใช้
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="math_problems.pdf", mimetype='application/pdf')


@app.route('/updateAnswer', methods=['POST'])
def updateAnswer():

    updated_problems = session.get('updated_problems', [])
    data = request.json
    idx = data.get('idx')
    txt = data.get('txt')

    if idx is not None and txt is not None:
        try:
            idx = int(idx)
            if 0 <= idx < len(updated_problems):
                # อัปเดตโจทย์ในตัวแปรเซสชัน
                updated_problems[idx] = txt
                session['updated_problems'] = updated_problems
                return jsonify({'success': True, 'message': 'Update successful'})
        except (ValueError, IndexError):
            pass

    return jsonify({'success': False, 'message': 'Invalid index or text'}), 400


@app.route('/get_words', methods=['GET'])
def get_words():
    word_type = request.args.get('type')
    words = []

    if word_type == 'noun':
        words = get_all_words("word_nouns1", "nouns", db2)  # ดึงคำนามทั้งหมด
    elif word_type == 'verb':
        words = get_all_words("word_verb", "verb", db2)
    elif word_type == 'object':
        words = get_all_words("word_object", "nouns", db2)  # ดึงกรรมทั้งหมด
    elif word_type == 'classifier':
        # ดึงลักษณะนามทั้งหมด
        words = get_all_words("word_classifier", "classifier", db2)

    return jsonify(words)


@app.route('/about')
def about():
    """
    flow = InstalledAppFlow.from_client_secrets_file(credentialsFile, SCOPES)
    creds = flow.run_local_server(port=0)
    print("token:", creds.to_json())
    """

    return render_template('about.html')

# ฟังก์ชันสำหรับหน้าแอดมินหลัก (แสดงทั้ง word_classifier และ mathprob)


@app.route('/admin', methods=['GET'])
def admin_page():
    if not session.get("is_logged_in"):
        return redirect(url_for('login'))

    if not session.get("isAdmin", False):
        return "คุณไม่มีสิทธิ์เข้าถึงหน้านี้", 403

    # ดึงข้อมูลลักษณะนาม
    classifiers = []
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor(dictionary=True)
        mycursor.execute("SELECT * FROM word_classifier")
        classifiers = mycursor.fetchall()
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาดในการเชื่อมต่อฐานข้อมูล: {err}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    # ดึงข้อมูลโจทย์คณิตศาสตร์
    problems = []
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db
        )
        mycursor = mydb.cursor(dictionary=True)
        mycursor.execute("SELECT * FROM mathprob_pattern")
        problems = mycursor.fetchall()
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาดในการเชื่อมต่อฐานข้อมูล: {err}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    nouns1 = []
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor(dictionary=True)
        mycursor.execute("SELECT * FROM word_nouns1")
        nouns1 = mycursor.fetchall()
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาดในการเชื่อมต่อฐานข้อมูล: {err}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

 # ดึงข้อมูลคำกริยา
    verbs = []
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor(dictionary=True)
        mycursor.execute("SELECT * FROM word_verb")
        verbs = mycursor.fetchall()
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาดในการเชื่อมต่อฐานข้อมูล: {err}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

# ดึงข้อมูลวัตถุ
    objects = []
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor(dictionary=True)
        mycursor.execute("SELECT * FROM word_object")
        objects = mycursor.fetchall()
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาดในการเชื่อมต่อฐานข้อมูล: {err}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return render_template('admin.html', classifiers=classifiers, problems=problems, nouns1=nouns1, verbs=verbs, objects=objects)

# ฟังก์ชันสำหรับเพิ่มข้อมูลใน word_classifier


@app.route('/admin/word_classifier', methods=['POST'])
def add_word_classifier():
    name = request.form.get('name')
    problem_type = request.form.get('type')
    mathsym = request.form.get('mathsym')

    if not name or not problem_type or not mathsym:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
    else:
        try:
            mydb = mysql.connector.connect(
                host=host, user=user, password=password, database=db2
            )
            mycursor = mydb.cursor()
            query = "INSERT INTO word_classifier (classifier, type, mathsym) VALUES (%s, %s, %s)"
            mycursor.execute(query, (name, problem_type, mathsym))
            mydb.commit()
            flash("ข้อมูลถูกเพิ่มเรียบร้อยแล้ว!", 'success')
        except mysql.connector.Error as err:
            flash(f"เกิดข้อผิดพลาด: {err}", 'error')
        finally:
            if 'mycursor' in locals():
                mycursor.close()
            if 'mydb' in locals():
                mydb.close()

    return redirect(url_for('admin_page'))


# ฟังก์ชันสำหรับลบข้อมูลใน word_classifier
@app.route('/admin/word_classifier/delete', methods=['POST'])
def delete_word_classifier():
    classifier_id = request.form['classifier_id']
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor()
        query = "DELETE FROM word_classifier WHERE id = %s"
        mycursor.execute(query, (classifier_id,))
        mydb.commit()
        flash("ลบข้อมูลสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {str(err)}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))

# ฟังก์ชันสำหรับแก้ไขข้อมูลใน word_classifier


@app.route('/admin/word_classifier/edit/<int:classifier_id>', methods=['POST'])
def edit_word_classifier(classifier_id):
    name = request.form.get('name')
    problem_type = request.form.get('type')
    mathsym = request.form.get('mathsym')

    if not name or not problem_type or not mathsym:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
        return redirect(url_for('admin_page'))

    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor()
        query = "UPDATE word_classifier SET classifier = %s, type = %s, mathsym = %s WHERE id = %s"
        mycursor.execute(query, (name, problem_type, mathsym, classifier_id))
        mydb.commit()
        flash("บันทึกการแก้ไขสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {str(err)}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))


# ฟังก์ชันสำหรับเพิ่มโจทย์คณิตศาสตร์ใหม่
@app.route('/admin/add_math_problem', methods=['POST'])
def add_math_problem():
    pattern = request.form.get('pattern')
    problem_type = request.form.get('type')
    mathsym = request.form.get('mathsym')

    if not pattern or not problem_type or not mathsym:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
    else:
        try:
            mydb = mysql.connector.connect(
                host=host, user=user, password=password, database=db
            )
            mycursor = mydb.cursor()
            query = "INSERT INTO mathprob_pattern (pattern, type, mathsym) VALUES (%s, %s, %s)"
            mycursor.execute(query, (pattern, problem_type, mathsym))
            mydb.commit()
            flash("โจทย์ถูกเพิ่มเรียบร้อยแล้ว!", 'success')
        except mysql.connector.Error as err:
            flash(f"เกิดข้อผิดพลาด: {err}", 'error')
        finally:
            if 'mycursor' in locals():
                mycursor.close()
            if 'mydb' in locals():
                mydb.close()

    return redirect(url_for('admin_page'))


# ฟังก์ชันสำหรับลบโจทย์คณิตศาสตร์
@app.route('/admin/delete_problem', methods=['POST'])
def delete_math_problem():
    problem_id = request.form['problem_id']
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db
        )
        mycursor = mydb.cursor()
        query = "DELETE FROM mathprob_pattern WHERE id = %s"
        mycursor.execute(query, (problem_id,))
        mydb.commit()
        flash("ลบโจทย์สำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {str(err)}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))


# ฟังก์ชันสำหรับแก้ไขโจทย์คณิตศาสตร์
@app.route('/admin/edit_problem/<int:problem_id>', methods=['POST'])
def edit_math_problem(problem_id):
    pattern = request.form.get('pattern')
    problem_type = request.form.get('type')
    mathsym = request.form.get('mathsym')

    if not pattern or not problem_type or not mathsym:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
        return redirect(url_for('admin_page'))

    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db
        )
        mycursor = mydb.cursor()
        query = "UPDATE mathprob_pattern SET pattern = %s, type = %s, mathsym = %s WHERE id = %s"
        mycursor.execute(query, (pattern, problem_type, mathsym, problem_id))
        mydb.commit()
        flash("บันทึกการแก้ไขสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {str(err)}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))

# ฟังก์ชันสำหรับเพิ่มข้อมูลใน word_nouns1


@app.route('/admin/word_nouns1', methods=['POST'])
def add_word_nouns1():
    nouns = request.form.get('nouns')  # ตรวจสอบชื่อให้ตรงกับฟอร์ม
    noun_type = request.form.get('type')  # ตรวจสอบชื่อให้ตรงกับฟอร์ม

    if not nouns or not noun_type:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
    else:
        try:
            mydb = mysql.connector.connect(
                host=host, user=user, password=password, database=db2
            )
            mycursor = mydb.cursor()
            # ชื่อคอลัมน์ในฐานข้อมูลต้องตรง
            query = "INSERT INTO word_nouns1 (nouns, type) VALUES (%s, %s)"
            mycursor.execute(query, (nouns, noun_type))
            mydb.commit()
            flash("ข้อมูลถูกเพิ่มเรียบร้อยแล้ว!", 'success')
        except mysql.connector.Error as err:
            flash(f"เกิดข้อผิดพลาด: {err}", 'error')
        finally:
            if 'mycursor' in locals():
                mycursor.close()
            if 'mydb' in locals():
                mydb.close()

    return redirect(url_for('admin_page'))  # ให้กลับไปที่หน้า admin


# ฟังก์ชันสำหรับลบข้อมูลใน word_nouns1
@app.route('/admin/word_nouns1/delete', methods=['POST'])
def delete_word_nouns1():
    noun_id = request.form['noun_id']
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor()
        query = "DELETE FROM word_nouns1 WHERE id = %s"
        mycursor.execute(query, (noun_id,))
        mydb.commit()
        flash("ลบข้อมูลสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {str(err)}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))


# ฟังก์ชันสำหรับแก้ไขข้อมูลใน word_nouns1
@app.route('/admin/word_nouns1/edit/<int:noun_id>', methods=['POST'])
def edit_word_nouns1(noun_id):
    nouns = request.form.get('nouns')  # เปลี่ยนชื่อให้ถูกต้อง
    noun_type = request.form.get('type')

    if not nouns or not noun_type:  # ใช้ nouns แทน noun
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
        return redirect(url_for('admin_page'))

    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor()
        # เปลี่ยนจาก noun เป็น nouns
        query = "UPDATE word_nouns1 SET nouns = %s, type = %s WHERE id = %s"
        mycursor.execute(query, (nouns, noun_type, noun_id)
                         )  # ใช้ nouns แทน noun
        mydb.commit()
        flash("บันทึกการแก้ไขสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {str(err)}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))


@app.route('/admin/word_verb', methods=['POST'])
def add_word_verb():
    verb = request.form.get('verb')
    mathsym = request.form.get('mathsym')

    if not verb or not mathsym:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
    else:
        try:
            mydb = mysql.connector.connect(
                host=host, user=user, password=password, database=db2
            )
            mycursor = mydb.cursor()
            query = "INSERT INTO word_verb (verb, mathsym) VALUES (%s, %s)"
            mycursor.execute(query, (verb, mathsym))
            mydb.commit()
            flash("ข้อมูลถูกเพิ่มเรียบร้อยแล้ว!", 'success')
        except mysql.connector.Error as err:
            flash(f"เกิดข้อผิดพลาด: {err}", 'error')
        finally:
            if 'mycursor' in locals():
                mycursor.close()
            if 'mydb' in locals():
                mydb.close()

    return redirect(url_for('admin_page'))

# ฟังก์ชันสำหรับลบข้อมูลใน word_verb


@app.route('/admin/word_verb/delete', methods=['POST'])
def delete_word_verb():
    verb_id = request.form['verb_id']
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor()
        query = "DELETE FROM word_verb WHERE id = %s"
        mycursor.execute(query, (verb_id,))
        mydb.commit()
        flash("ลบข้อมูลสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {err}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))

# ฟังก์ชันสำหรับแก้ไขข้อมูลใน word_verb


@app.route('/admin/word_verb/edit/<int:verb_id>', methods=['POST'])
def edit_word_verb(verb_id):
    verb = request.form.get('verb')
    mathsym = request.form.get('mathsym')

    if not verb or not mathsym:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
        return redirect(url_for('admin_page'))

    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor()
        query = "UPDATE word_verb SET verb = %s, mathsym = %s WHERE id = %s"
        mycursor.execute(query, (verb, mathsym, verb_id))
        mydb.commit()
        flash("บันทึกการแก้ไขสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {err}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))

# ฟังก์ชันสำหรับเพิ่มข้อมูลใน word_object


@app.route('/admin/word_object', methods=['POST'])
def add_word_object():
    nouns = request.form.get('nouns')  # เปลี่ยนจาก noun เป็น nouns
    object_type = request.form.get('type')  # ใช้ object_type

    if not nouns or not object_type:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
    else:
        try:
            mydb = mysql.connector.connect(
                host=host, user=user, password=password, database=db2
            )
            mycursor = mydb.cursor()
            # เปลี่ยนจาก noun เป็น nouns
            query = "INSERT INTO word_object (nouns, type) VALUES (%s, %s)"
            mycursor.execute(query, (nouns, object_type))
            mydb.commit()
            flash("ข้อมูลถูกเพิ่มเรียบร้อยแล้ว!", 'success')
        except mysql.connector.Error as err:
            flash(f"เกิดข้อผิดพลาด: {err}", 'error')
        finally:
            if 'mycursor' in locals():
                mycursor.close()
            if 'mydb' in locals():
                mydb.close()

    return redirect(url_for('admin_page'))


# ฟังก์ชันสำหรับลบข้อมูลใน word_object
@app.route('/admin/word_object/delete', methods=['POST'])
def delete_word_object():
    object_id = request.form['object_id']
    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor()
        query = "DELETE FROM word_object WHERE id = %s"
        mycursor.execute(query, (object_id,))
        mydb.commit()
        flash("ลบข้อมูลสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {str(err)}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))


# ฟังก์ชันสำหรับแก้ไขข้อมูลใน word_object
@app.route('/admin/word_object/edit/<int:object_id>', methods=['POST'])
def edit_word_object(object_id):
    nouns = request.form.get('nouns')
    type_ = request.form.get('type')

    if not nouns or not type_:
        flash("กรุณากรอกข้อมูลให้ครบถ้วน", 'error')
        return redirect(url_for('admin_page'))

    try:
        mydb = mysql.connector.connect(
            host=host, user=user, password=password, database=db2
        )
        mycursor = mydb.cursor()
        query = "UPDATE word_object SET nouns = %s, type = %s WHERE id = %s"
        mycursor.execute(query, (nouns, type_, object_id))
        mydb.commit()
        flash("บันทึกการแก้ไขสำเร็จ!", 'success')
    except mysql.connector.Error as err:
        flash(f"เกิดข้อผิดพลาด: {str(err)}", 'error')
    finally:
        if 'mycursor' in locals():
            mycursor.close()
        if 'mydb' in locals():
            mydb.close()

    return redirect(url_for('admin_page'))


# หน้า home
@app.route("/")
def home_redirect():
    return redirect(url_for('index'))  # เปลี่ยนเส้นทางไปที่หน้าหลัก /home

# หน้า login


@app.route("/login")
def login():
    return render_template("login.html")

# หน้าsignup


@app.route("/signup")
def signup():
    return render_template("signup.html")

# เช็คความเเข็งเเรงของpassword


def check_password_strength(password):
    # รหัสต้องประกอบไปด้วยตัวพิมพ์เล็ก,พิมพ์ใหญ่,ตัวเลข,สัญลักษณ์พิเศษ อย่างน้อย8ตัวอักษร
    return re.match(r'^(?=.*\d)(?=.*[!@#$%^&*])(?=.*[a-z])(?=.*[A-Z]).{8,}$', password) is not None

# ยืนยันตัวตนเเละจัดการสิทธ์(ผู้ใช้ทั่วไปเเละadmin)


@app.route("/result", methods=["POST", "GET"])
def result():
    if request.method == "POST":
        result = request.form
        email = result["email"]
        password = result["pass"]
        try:
            # ยืนยันข้อมูลlogin
            user = auth.sign_in_with_email_and_password(email, password)
            session["is_logged_in"] = True
            session["email"] = user["email"]
            session["uid"] = user["localId"]

            # ดึงข้อมูลมาตรวจสอบ
            data = dbFirebase.child("users").child(session["uid"]).get().val()
            print("Firebase data:", data)

            # ตรวจสอบการเป็นadmin
            if data and "isAdmin" in data:
                session["isAdmin"] = data["isAdmin"]
                print("session isAdmin:", session["isAdmin"])
            else:
                # กรณีไม่พบ isAdmin, ตั้งค่าเป็น False
                session["isAdmin"] = False

            # เข้าสู่หน้าต่างๆตามประเภทผู้ใช้
            if session["isAdmin"]:
                return redirect(url_for('admin_page'))
            else:
                return redirect(url_for('index'))

        except Exception as e:
            print("Error occurred:", e)
            return redirect(url_for('login'))
    else:
        if session.get("is_logged_in", False):
            return redirect(url_for('index'))
        else:
            return redirect(url_for('login'))


# ส่วนของหน้าสมัครใช้งาน
@app.route("/register", methods=["POST", "GET"])
def register():
    if request.method == "POST":
        result = request.form
        email = result["email"]
        password = result["pass"]
        name = result["name"]
        if not check_password_strength(password):
            print("Password does not meet strength requirements")
            return redirect(url_for('signup'))
        try:
            # สร้างบัญชีผู้ใช้
            auth.create_user_with_email_and_password(email, password)
            # จัดเก็บข้อมูล
            user = auth.sign_in_with_email_and_password(email, password)
            session["is_logged_in"] = True
            session["email"] = user["email"]
            session["uid"] = user["localId"]
            session["name"] = name
            # บันทึกข้อมูลผู้ใช้ลง firebase
            data = {"name": name, "email": email,
                    "last_logged_in": datetime.now().strftime("%m/%d/%Y, %H:%M:%S")}
            dbFirebase.child("users").child(session["uid"]).set(data)
            return redirect(url_for('index'))
        except Exception as e:
            print("Error occurred during registration: ", e)
            return redirect(url_for('signup'))
    else:
        # เข้าสู่หน้าเเรกหลังจาก login
        if session.get("is_logged_in", False):
            return redirect(url_for('index'))
        else:
            return redirect(url_for('signup'))

# ส่วนของรีเซตpassword


@app.route("/reset_password", methods=["GET", "POST"])
def reset_password():
    if request.method == "POST":
        email = request.form["email"]
        try:
            # Send password reset email
            auth.send_password_reset_email(email)
            return render_template("reset_password_done.html")
        except Exception as e:
            print("Error occurred: ", e)
            return render_template("reset_password.html", error="An error occurred. Please try again.")
    else:
        return render_template("reset_password.html")

# ส่วนของการ logout


@app.route("/logout")
def logout():
    # เวลาที่ออกจากระบบ
    if "uid" in session:
        dbFirebase.child("users").child(session["uid"]).update(
            {"last_logged_out": datetime.now().strftime("%m/%d/%Y, %H:%M:%S")})

    # คืนค่าต่างๆ
    session["is_logged_in"] = False
    session["isAdmin"] = False

    # เปลี่ยนเส้นทางไปที่หน้า home แทน login
    return redirect(url_for('index'))

# สร้างPDFของส่วนประวัติการใช้งาน


@app.route('/generate_selected_pdf', methods=['POST'])
def generate_selected_pdf():
    selected_indices = request.form.getlist('selected_problems')

    updated_problems = session.get('updated_problems', [])

    # ตรวจสอบว่า index ที่ส่งมาถูกต้องหรือไม่
    selected_problems = []
    for idx in selected_indices:
        try:
            selected_problems.append(updated_problems[int(idx)])
        except IndexError:
            print(f"Index {idx} out of range, skipping...")
            continue

    if not selected_problems:
        return "No problems selected or invalid indices", 400

    # สร้างไฟล์ PDF ต่อไปเหมือนเดิม
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)

    # เพิ่มฟอนต์ภาษาไทย
    pdfmetrics.registerFont(TTFont('THSarabun', 'THSarabunNew.ttf'))

    # ตั้งค่าฟอนต์และขนาด
    text_object = c.beginText(50, 750)
    text_object.setFont("THSarabun", 16)

    # สร้างหัวกระดาษ ชื่อ,นามสกุล,เลขที่,วันที่
    text_object.textLine(
        f"ชื่อ: __________  นามสกุล: __________  เลขที่: ____        วันที่: {datetime.now().strftime('%d/%m/%Y')}")
    c.drawText(text_object)

    c.setFont("THSarabun", 16)
    c.drawRightString(580, 750, "คะแนน: ____")

    # เขียนหัวกระดาษ
    c.setFont("THSarabun", 20)
    c.drawCentredString(300, 720, "โจทย์ปัญหาคณิตศาสตร์")

    # ปรับส่วนต่างๆของกระดาษ
    text_object = c.beginText(50, 680)
    text_object.setFont("THSarabun", 16)

    for i, problem in enumerate(selected_problems, 1):
        text_object.textLine(f"{i}) {problem}")
        text_object.moveCursor(0, 10)
        text_object.textLine(
            "ตอบ: .....................................................................................")
        text_object.moveCursor(0, 5)

        if text_object.getY() < 100:
            c.drawText(text_object)
            c.showPage()
            text_object = c.beginText(50, 750)
            text_object.setFont("THSarabun", 16)

    c.drawText(text_object)

    for page_num in range(1, c.getPageNumber() + 1):
        c.setFont("THSarabun", 12)
        c.drawRightString(580, 20, f"Page {page_num}")

    c.save()

    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="selected_problems.pdf", mimetype='application/pdf')


# ดาวโหลด .dock

@app.route('/download_docx', methods=['POST'])
def download_docx():
    updated_problems = session.get('updated_problems', [])

    # ตรวจสอบว่าโจทย์ไม่ว่างเปล่า
    if not updated_problems:
        return "No problems to download", 400

    # สร้างไฟล์ DOCX
    doc = Document()  # type: ignore

    # ตั้งค่าฟอนต์เริ่มต้นเป็นฟอนต์ภาษาไทย
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)

    # สร้างหัวกระดาษ
    header_text = f"ชื่อ: ____________  นามสกุล: ____________  เลขที่: ______        วันที่: {datetime.now().strftime('%d/%m/%Y')}"
    header_paragraph = doc.add_paragraph(header_text)
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # เพิ่มช่องคะแนนที่มุมบนขวาของเอกสาร
    score_paragraph = doc.add_paragraph("คะแนน: ______")
    score_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # เพิ่มหัวข้อของเอกสาร
    title = doc.add_heading('โจทย์ปัญหาคณิตศาสตร์', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # เพิ่มโจทย์และบรรทัดคำตอบ
    for i, problem in enumerate(updated_problems, 1):
        # ทำความสะอาด HTML tags ก่อน
        clean_problem = clean_html_tags(problem)
        doc.add_paragraph(f"{i}) {clean_problem}")
        answer_line = doc.add_paragraph(
            "ตอบ: .....................................................................................")
        answer_line_format = answer_line.paragraph_format
        answer_line_format.space_after = Pt(12)  # เพิ่มระยะห่างระหว่างบรรทัด

    # บันทึกไฟล์ลงใน memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # ส่งไฟล์ DOCX กลับให้ผู้ใช้
    return send_file(buffer, as_attachment=True, download_name="math_problems.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# บันทึกโจทย์เป็นไฟล์ word ในหน้าการบันทึก


def clean_html_tags(html_content):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, "html.parser")
    return soup.get_text()


@app.route('/d', methods=['POST'])
def d():
    # รับ set_idx จาก request
    set_idx = int(request.form.get('set_idx'))  # เปลี่ยนเป็น form.get
    uid = session["uid"]

    # ดึงชุดโจทย์จาก Firebase
    data = dbFirebase.child("problems").child(uid).get().val()
    if not data or 'sets' not in data or set_idx >= len(data['sets']):
        return "No problems to download", 400

    problem_set = data['sets'][set_idx]
    problems = problem_set["problems"]
    answers = problem_set["answers"]
    units = problem_set["units"]

    # สร้างไฟล์ DOCX
    doc = Document()

    # ตั้งค่าฟอนต์เริ่มต้นเป็นฟอนต์ภาษาไทย
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)

    # สร้างหัวกระดาษ ชื่อ,นามสกุล,เลขที่,วันที่
    header_text = f"ชื่อ: ____________  นามสกุล: ____________  เลขที่: ______        วันที่: {datetime.now().strftime('%d/%m/%Y')}"
    header_paragraph = doc.add_paragraph(header_text)
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # เพิ่มช่องคะแนนที่มุมบนขวาของเอกสาร
    score_paragraph = doc.add_paragraph("คะแนน: ______")
    score_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # เพิ่มหัวข้อของเอกสาร
    title = doc.add_heading('โจทย์ปัญหาคณิตศาสตร์', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # เพิ่มโจทย์และบรรทัดคำตอบ
    for i, problem in enumerate(problems, 1):
        # ทำความสะอาด HTML tags ก่อนเขียนลง DOCX
        clean_problem = clean_html_tags(problem)
        doc.add_paragraph(f"{i}) {clean_problem}")

        # ตรวจสอบว่ามีคำตอบและหน่วยตรงตามดัชนี i หรือไม่
        if i - 1 < len(answers) and i - 1 < len(units):
            answer_text = f"ตอบ: {answers[i - 1]} {units[i - 1]}"
        else:
            # ถ้าไม่มีคำตอบหรือหน่วยในดัชนีนั้นให้ตั้งค่าเป็นว่าง
            answer_text = "ตอบ: ______________"

        answer_line = doc.add_paragraph(answer_text)
        answer_line.paragraph_format.space_after = Pt(
            12)  # เพิ่มระยะห่างระหว่างบรรทัด

    # บันทึกไฟล์ลงใน memory และส่งให้ผู้ใช้
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name=f"problem_set_{set_idx + 1}.docx",
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# ปุ่มลบบันทึกโจทย์ที่บันทึกเป็นชุดๆ
@app.route("/delete_set/<int:set_idx>", methods=["POST"])
def delete_set(set_idx):
    # ตรวจสอบว่าผู้ใช้ได้ล็อกอินหรือไม่
    if not session.get("is_logged_in"):
        return redirect(url_for('login'))

    # ดึงข้อมูลผู้ใช้ที่ล็อกอินอยู่
    uid = session.get("uid")

    # ตรวจสอบว่ามีข้อมูลชุดโจทย์ใน Firebase
    data = dbFirebase.child("problems").child(uid).get().val()
    if not data or 'sets' not in data or set_idx >= len(data['sets']):
        return "ไม่พบชุดโจทย์ที่ต้องการลบ", 404

    # ลบชุดโจทย์ที่ระบุ
    data['sets'].pop(set_idx)

    # อัปเดตข้อมูลใน Firebase
    dbFirebase.child("problems").child(uid).set(data)

    return redirect(url_for("show_all"))


# ทำความสะอาดข้อความโจทย์ก่อนที่จะนำไปบันทึกใน PDF หรือ DOCX


def clean_html_tags(text):
    # ลบ HTML tags ทั้งหมด
    clean_text = re.sub(r'<.*?>', '', text)
    return clean_text


if __name__ == "__main__":
    app.run(debug=True)
